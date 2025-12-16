import pymysql
from docx import Document
import re

# -------------------------- 1. 配置项（必须替换为你的信息） --------------------------
# MySQL连接配置
MYSQL_CONFIG = {
    "host": "127.0.0.1",
    "port": 3306,
    "user": "root",  # 你的MySQL用户名
    "password": "123456",  # 如Root@123456
    "db": "driving_exam_system",
    "charset": "utf8mb4"
}

# 题库文件路径（替换为你的两个题库路径）
SUBJECT1_FILE = "C:\\Users\\Lenovo\\Desktop\\C1C2科目一题库.docx"
SUBJECT4_FILE = "C:\\Users\\Lenovo\\Desktop\\C1C2科目四题库.docx"

# 题型映射（对应question_type表的type_id）
QUESTION_TYPE_MAP = {
    "单选": 1,
    "多选": 2,
    "判断": 3
}

# -------------------------- 2. 数据库工具函数 --------------------------
def get_db_conn():
    """创建MySQL连接，处理异常"""
    try:
        conn = pymysql.connect(**MYSQL_CONFIG)
        return conn
    except pymysql.MySQLError as e:
        print(f"❌ 数据库连接失败：{e}")
        exit(1)

def batch_insert_questions(questions):
    """批量插入题目到question_bank表"""
    if not questions:
        print("❌ 无有效题目，跳过插入")
        return

    conn = get_db_conn()
    cursor = conn.cursor()
    # 插入SQL（匹配question_bank表字段）
    insert_sql = """
    INSERT INTO question_bank (
        subject_type, type_id, question_content, option_a, option_b,
        option_c, option_d, correct_answer, analysis, score,
        difficulty, has_image, image_path
    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """

    # 构造插入数据（处理空值）
    insert_data = []
    for q in questions:
        data = (
            q["subject_type"], q["type_id"], q["question_content"],
            q["option_a"] or "", q["option_b"] or "", q["option_c"] or "", q["option_d"] or "",
            q["correct_answer"], q["analysis"] or "", q["score"] or 1,
            q["difficulty"] or "易", q["has_image"], q["image_path"] or ""
        )
        insert_data.append(data)

    try:
        cursor.executemany(insert_sql, insert_data)
        conn.commit()
        print(f"✅ 成功插入 {cursor.rowcount} 道题目（{q['subject_type']}）")
    except pymysql.MySQLError as e:
        conn.rollback()
        print(f"❌ 插入失败：{e}")
    finally:
        cursor.close()
        conn.close()

# -------------------------- 3. 解析题库文件（核心逻辑） --------------------------
def parse_subject1_docx(file_path):
    """解析科目一题库（docx格式）：提取单选/判断，处理图片标记"""
    doc = Document(file_path)
    questions = []
    current_q = {}  # 临时存储当前题目
    option_pattern = re.compile(r"^[A-D]、")  # 匹配选项（如"A、"）
    answer_pattern = re.compile(r"答案：([A-Z√×]+)")  # 匹配答案（如"答案：D"）
    image_pattern = re.compile(r"!\[img\]\([^)]+\)")  # 匹配图片标记

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 1. 识别题目编号（如"1.机动车驾驶人饮酒后..."），开启新题目
        if re.match(r"^\d+[.、]", text):
            # 若有未保存的上一题，先加入列表
            if current_q:
                # 补充科目一默认信息
                current_q.update({
                    "subject_type": "科目一",
                    "score": 1,
                    "difficulty": "易",
                    "analysis": "",
                    "has_image": 1 if current_q.get("image_path") else 0
                })
                # 区分题型：判断（答案为√/×）、单选（答案为A-D）
                if current_q["correct_answer"] in ["√", "×"]:
                    current_q["type_id"] = QUESTION_TYPE_MAP["判断"]
                else:
                    current_q["type_id"] = QUESTION_TYPE_MAP["单选"]
                questions.append(current_q)
            # 初始化新题目
            current_q = {
                "question_content": text,
                "option_a": "", "option_b": "", "option_c": "", "option_d": "",
                "correct_answer": "", "image_path": ""
            }
            # 检查题干是否含图片标记
            image_match = image_pattern.search(text)
            if image_match:
                current_q["image_path"] = image_match.group()  # 保存图片标记
                current_q["question_content"] = image_pattern.sub("", text)  # 移除图片标记，保留纯题干

        # 2. 识别选项（A、B、C、D）
        elif option_pattern.match(text):
            option_key = "option_" + text[0].lower()  # 转为option_a/option_b
            option_content = text[2:].strip()  # 提取选项内容（如"A、1分"→"1分"）
            current_q[option_key] = option_content

        # 3. 识别答案
        elif answer_pattern.search(text):
            current_q["correct_answer"] = answer_pattern.search(text).group(1)

    # 保存最后一道题
    if current_q:
        current_q.update({
            "subject_type": "科目一", "score": 1, "difficulty": "易", "analysis": "",
            "has_image": 1 if current_q.get("image_path") else 0,
            "type_id": QUESTION_TYPE_MAP["判断"] if current_q["correct_answer"] in ["√", "×"] else QUESTION_TYPE_MAP["单选"]
        })
        questions.append(current_q)

    print(f"📊 解析科目一题库完成，共提取 {len(questions)} 道题目")
    return questions

def parse_subject4_docx(file_path):
    """解析科目四题库（docx格式）：提取单选/多选，处理情景题"""
    doc = Document(file_path)
    questions = []
    current_q = {}
    option_pattern = re.compile(r"^[A-D]、")
    answer_pattern = re.compile(r"答案：([A-Z,]+|正确|错误)")  # 多选答案可能为"AB"，判断为"正确/错误"
    image_pattern = re.compile(r"!\[img\]\([^)]+\)")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 1. 识别题目编号（如"1. 机动车在路边起步后..."）
        if re.match(r"^\d+[.、]", text):
            if current_q:
                # 补充科目四默认信息
                current_q.update({
                    "subject_type": "科目四",
                    "score": 1,
                    "difficulty": "易",
                    "analysis": "",
                    "has_image": 1 if current_q.get("image_path") else 0
                })
                # 区分题型：多选（答案含多个字母，如"AB"）、单选（单个字母）、判断（正确/错误）
                if len(current_q["correct_answer"]) > 1 and current_q["correct_answer"] not in ["正确", "错误"]:
                    current_q["type_id"] = QUESTION_TYPE_MAP["多选"]
                elif current_q["correct_answer"] in ["正确", "错误"]:
                    current_q["type_id"] = QUESTION_TYPE_MAP["判断"]
                else:
                    current_q["type_id"] = QUESTION_TYPE_MAP["单选"]
                questions.append(current_q)
            # 初始化新题目
            current_q = {
                "question_content": text,
                "option_a": "", "option_b": "", "option_c": "", "option_d": "",
                "correct_answer": "", "image_path": ""
            }
            # 处理图片标记
            image_match = image_pattern.search(text)
            if image_match:
                current_q["image_path"] = image_match.group()
                current_q["question_content"] = image_pattern.sub("", text)

        # 2. 识别选项
        elif option_pattern.match(text):
            option_key = "option_" + text[0].lower()
            current_q[option_key] = text[2:].strip()

        # 3. 识别答案（处理"正确/错误"转为"√/×"，统一格式）
        elif answer_pattern.search(text):
            ans = answer_pattern.search(text).group(1)
            current_q["correct_answer"] = "√" if ans == "正确" else ("×" if ans == "错误" else ans)

    # 保存最后一道题
    if current_q:
        current_q.update({
            "subject_type": "科目四", "score": 1, "difficulty": "易", "analysis": "",
            "has_image": 1 if current_q.get("image_path") else 0,
            "type_id": QUESTION_TYPE_MAP["多选"] if (len(current_q["correct_answer"]) > 1 and current_q["correct_answer"] not in ["√", "×"]) else QUESTION_TYPE_MAP["单选"]
        })
        questions.append(current_q)

    print(f"📊 解析科目四题库完成，共提取 {len(questions)} 道题目")
    return questions

# -------------------------- 4. 数据验证函数 --------------------------
def verify_data():
    """验证插入结果：统计各科目/题型数量+随机抽查3道题"""
    conn = get_db_conn()
    cursor = conn.cursor(pymysql.cursors.DictCursor)

    # 1. 按科目+题型统计
    print("\n===== 数据统计结果 =====")
    cursor.execute("""
        SELECT q.subject_type, t.type_name, COUNT(q.question_id) AS total
        FROM question_bank q
        LEFT JOIN question_type t ON q.type_id = t.type_id
        GROUP BY q.subject_type, q.type_id
        ORDER BY q.subject_type, q.type_id
    """)
    stats = cursor.fetchall()
    for stat in stats:
        print(f"科目：{stat['subject_type']} | 题型：{stat['type_name']} | 题目数量：{stat['total']}")

    # 2. 随机抽查3道题（含图片的题目优先）
    print("\n===== 随机抽查3道题 =====")
    cursor.execute("""
        SELECT q.question_id, q.subject_type, t.type_name, q.question_content, q.correct_answer, q.has_image
        FROM question_bank q
        LEFT JOIN question_type t ON q.type_id = t.type_id
        ORDER BY RAND() LIMIT 3
    """)
    samples = cursor.fetchall()
    for i, sample in enumerate(samples, 1):
        print(f"\n第{i}道【{sample['subject_type']}-{sample['type_name']}】")
        print(f"题目ID：{sample['question_id']}")
        print(f"题干：{sample['question_content'][:80]}..." if len(sample['question_content'])>80 else f"题干：{sample['question_content']}")
        print(f"正确答案：{sample['correct_answer']} | 有无图片：{'有' if sample['has_image']==1 else '无'}")

    cursor.close()
    conn.close()

# -------------------------- 5. 主函数：执行解析+插入+验证 --------------------------
if __name__ == "__main__":
    print("=== 开始解析C1C2驾照题库 ===")
    # 步骤1：解析科目一题库并插入
    subject1_questions = parse_subject1_docx(SUBJECT1_FILE)
    batch_insert_questions(subject1_questions)
    # 步骤2：解析科目四题库并插入
    subject4_questions = parse_subject4_docx(SUBJECT4_FILE)
    batch_insert_questions(subject4_questions)
    # 步骤3：验证数据
    verify_data()
    print("\n=== 题库填充+验证完成 ===")